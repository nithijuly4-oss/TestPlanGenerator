#!/usr/bin/env python3
"""
Tool: create_docx_plan.py
Purpose: Create DOCX test plan document with content based on LLM-generated test plan
Usage: python tools/create_docx_plan.py <ISSUE_KEY> <ISSUE_TITLE> <JSON_FILE_PATH> [TEMP_DIR]
"""

import os
import sys
import json
import uuid
import zipfile
import shutil
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

# Ensure UTF-8 encoding for stdout (Windows compatibility)
if sys.version_info >= (3, 7):
    try:
        sys.stdout.reconfigure(encoding='utf-8')
    except Exception:
        pass

# Get temp directory from command line or use default
TEMP_DIR = sys.argv[4] if len(sys.argv) > 4 else '.tmp'

def remove_headers_footers_from_docx(docx_path):
    """Remove all headers, footers, and watermarks by stripping from DOCX XML."""
    
    # Create temp directory
    temp_dir = os.path.join(TEMP_DIR, f"_docx_temp_{uuid.uuid4().hex[:8]}")
    os.makedirs(temp_dir, exist_ok=True)
    
    try:
        # Extract DOCX
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            zip_ref.extractall(temp_dir)
        
        # Step 1: Clean document.xml to remove header/footer references
        doc_xml_path = os.path.join(temp_dir, 'word', 'document.xml')
        if os.path.exists(doc_xml_path):
            with open(doc_xml_path, 'r', encoding='utf-8') as f:
                doc_xml = f.read()
            
            # Remove header/footer references from section properties
            doc_xml = re.sub(r'<w:headerReference[^>]*?/?>', '', doc_xml)
            doc_xml = re.sub(r'<w:footerReference[^>]*?/?>', '', doc_xml)
            
            with open(doc_xml_path, 'w', encoding='utf-8') as f:
                f.write(doc_xml)
        
        # Step 2: Clean document.xml.rels to remove header/footer relationships
        rels_path = os.path.join(temp_dir, 'word', '_rels', 'document.xml.rels')
        if os.path.exists(rels_path):
            with open(rels_path, 'r', encoding='utf-8') as f:
                rels_xml = f.read()
            
            # Remove relationships to header and footer files
            rels_xml = re.sub(r'<Relationship[^>]*Target="header\d+\.xml"[^>]*/?>', '', rels_xml)
            rels_xml = re.sub(r'<Relationship[^>]*Target="footer\d+\.xml"[^>]*/?>', '', rels_xml)
            
            with open(rels_path, 'w', encoding='utf-8') as f:
                f.write(rels_xml)
        
        # Step 3: Rebuild DOCX, filtering out header and footer files
        with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
            for root, dirs, files in os.walk(temp_dir):
                for file in files:
                    # Skip header and footer XML files
                    if (file.startswith('header') or file.startswith('footer')) and file.endswith('.xml'):
                        continue
                    
                    file_path = os.path.join(root, file)
                    arcname = os.path.relpath(file_path, temp_dir)
                    zip_ref.write(file_path, arcname)
    
    finally:
        # Clean up temp directory
        shutil.rmtree(temp_dir, ignore_errors=True)


def create_docx_plan(issue_key, issue_title, test_plan_sections):
    """Create DOCX file using template styling but with Jira-specific content only."""
    
    try:
        # Step 1: Load template for styles/formatting only
        template_path = "./test_plan_templat/Test Plan - Template.docx"
        
        if not os.path.exists(template_path):
            return {
                "status": "error",
                "message": f"Template not found at {template_path}",
                "timestamp": datetime.now().isoformat()
            }
        
        doc = Document(template_path)
        
        # Step 2: Remove all headers, footers, and watermarks completely
        for section in doc.sections:
            # Remove entire header section
            header = section.header
            # Delete all elements from header
            for element in header._element:
                header._element.remove(element)
            
            # Remove entire footer section  
            footer = section.footer
            # Delete all elements from footer
            for element in footer._element:
                footer._element.remove(element)
            
            # Remove any watermarks from document properties
            try:
                # Clear any background/watermark elements
                for para in header.paragraphs:
                    para.clear()
                for para in footer.paragraphs:
                    para.clear()
            except:
                pass
        
        # Step 3: Delete all original template content (all paragraphs and tables)
        # except we'll keep the document structure for styles
        elements_to_delete = []
        para_count = 0
        for element in doc.element.body:
            if element.tag.endswith('}p'):  # Paragraph
                para_count += 1
                if para_count > 0:  # Delete all paragraphs
                    elements_to_delete.append(element)
            elif element.tag.endswith('}tbl'):  # Table
                elements_to_delete.append(element)
        
        # Delete elements in reverse order
        for element in reversed(elements_to_delete):
            element.getparent().remove(element)
        
        # Step 4: Add fresh title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(f"Test Plan - {issue_key}: {issue_title}")
        title_run.bold = True
        title_run.font.size = Pt(16)
        title_para.style = 'Heading 1'
        
        # Step 5: Define required sections with Jira-specific content
        required_sections = [
            ("Objective", test_plan_sections.get("objective", "Test objectives to be defined.")),
            ("Scope", test_plan_sections.get("scope", "Test scope to be defined.")),
            ("Inclusions", test_plan_sections.get("inclusions", "Items to be included in testing to be defined.")),
            ("Test Environments", test_plan_sections.get("test_environments", "Testing environments to be defined.")),
            ("Defect Reporting Procedure", test_plan_sections.get("defect_reporting", "Defects will be logged in JIRA with detailed steps for reproduction.")),
            ("Test Strategy", test_plan_sections.get("test_strategy", "Testing strategy to be defined.")),
            ("Test Schedule", test_plan_sections.get("test_schedule", "Testing will be performed in multiple cycles based on project schedule.")),
            ("Test Deliverables", test_plan_sections.get("test_deliverables", "Test delivery items will include test cases, reports, and defect logs.")),
        ]
        
        # Step 6: Build document with required sections
        for section_name, content in required_sections:
            # Add section heading (Level 2)
            heading_para = doc.add_heading(section_name, level=2)
            heading_para.style = 'Heading 2'
            
            # Add section content - handle both string and list content
            content_para = doc.add_paragraph()
            if content:
                # Convert to string if it's a list
                if isinstance(content, list):
                    content_str = " ".join([str(item) for item in content if item])
                else:
                    content_str = str(content).strip() if content else ""
                
                if content_str:
                    content_para.add_run(content_str)
                else:
                    content_para.add_run(f"{section_name} details not provided.")
            else:
                content_para.add_run(f"{section_name} details not provided.")
        
        # Step 7: Add optional sections - handle both string and list content
        if "entry_and_exit" in test_plan_sections and test_plan_sections["entry_and_exit"]:
            content = test_plan_sections["entry_and_exit"]
            if isinstance(content, list):
                content = " ".join([str(item) for item in content if item])
            doc.add_heading("Entry and Exit Criteria", level=2)
            doc.add_paragraph(str(content))
        
        if "risks" in test_plan_sections and test_plan_sections["risks"]:
            content = test_plan_sections["risks"]
            if isinstance(content, list):
                content = " ".join([str(item) for item in content if item])
            doc.add_heading("Risks and Mitigations", level=2)
            doc.add_paragraph(str(content))
        
        if "tools" in test_plan_sections and test_plan_sections["tools"]:
            content = test_plan_sections["tools"]
            if isinstance(content, list):
                content = " ".join([str(item) for item in content if item])
            doc.add_heading("Tools", level=2)
            doc.add_paragraph(str(content))
        
        # Step 8: Add Test Cases section with table
        test_cases = test_plan_sections.get("test_cases", [])
        if test_cases:
            doc.add_heading("Test Cases", level=2)
            
            # Create test cases table
            table = doc.add_table(rows=1, cols=4)
            try:
                # Try to apply table style from template
                table.style = 'Light Grid Accent 1'
            except:
                pass  # Use default if style not available
            
            # Set header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Test ID'
            hdr_cells[1].text = 'Test Case Title'
            hdr_cells[2].text = 'Test Steps'
            hdr_cells[3].text = 'Expected Result'
            
            # Add test case rows
            for tc in test_cases[:15]:  # Limit to 15 test cases
                row_cells = table.add_row().cells
                row_cells[0].text = tc.get("id", "")
                row_cells[1].text = str(tc.get("title", ""))[:100]
                row_cells[2].text = str(tc.get("steps", ""))[:150]
                row_cells[3].text = str(tc.get("expected_result", ""))[:100]
        
        # Step 9: Save document
        os.makedirs(TEMP_DIR, exist_ok=True)
        file_id = str(uuid.uuid4())[:8]
        output_path = os.path.join(TEMP_DIR, f"{file_id}.docx")
        doc.save(output_path)
        
        # Step 10: Remove any remaining headers, footers, and watermarks
        try:
            remove_headers_footers_from_docx(output_path)
        except Exception as e:
            pass  # Silently fail if removal has issues
        
        # Step 11: Verify file was created
        if os.path.exists(output_path):
            file_size_kb = os.path.getsize(output_path) / 1024
            return {
                "status": "success",
                "file_path": output_path,
                "file_name": os.path.basename(output_path),
                "file_id": file_id,
                "issue_key": issue_key,
                "file_size_kb": round(file_size_kb, 2),
                "timestamp": datetime.now().isoformat()
            }
        else:
            return {
                "status": "error",
                "message": "File was not created successfully",
                "timestamp": datetime.now().isoformat()
            }
    
    except Exception as e:
        return {
            "status": "error",
            "message": f"Error creating DOCX: {str(e)[:100]}",
            "timestamp": datetime.now().isoformat()
        }

if __name__ == "__main__":
    if len(sys.argv) < 3:
        sys.stderr.write("Usage: python create_docx_plan.py <ISSUE_KEY> <ISSUE_TITLE> [JSON_FILE_PATH or JSON_STRING]\n")
        sys.exit(1)
    
    issue_key = sys.argv[1]
    issue_title = sys.argv[2]
    json_input = sys.argv[3] if len(sys.argv) > 3 else None
    
    # Load test plan - try file first, then JSON string
    test_plan = {}
    if json_input:
        try:
            # Try as file path first
            if os.path.exists(json_input):
                with open(json_input, 'r', encoding='utf-8') as f:
                    test_plan = json.load(f)
            else:
                # Try as JSON string
                try:
                    test_plan = json.loads(json_input)
                except json.JSONDecodeError:
                    sys.stderr.write(f"Warning: Could not parse as file or JSON: {json_input[:50]}\n")
        except Exception as e:
            sys.stderr.write(f"Error reading JSON: {str(e)}\n")
            sys.exit(1)
    
    result = create_docx_plan(issue_key, issue_title, test_plan)
    
    # Output JSON only (no debug prints)
    print(json.dumps(result, ensure_ascii=True))
