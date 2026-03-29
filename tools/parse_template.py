#!/usr/bin/env python3
"""
Tool: parse_template.py
Purpose: Parse DOCX template following SOP-003
Usage: python tools/parse_template.py
"""

import os
import json
from datetime import datetime

def parse_template():
    """Parse test plan template DOCX file."""
    
    print("🔍 Parsing Test Plan Template...")
    print("-" * 50)
    
    try:
        # Step 1: Check if file exists
        template_path = "./test_plan_templat/Test Plan - Template.docx"
        
        if not os.path.exists(template_path):
            return {
                "status": "error",
                "template_name": "Test Plan - Template.docx",
                "message": f"Template file not found at: {template_path}",
                "timestamp": datetime.now().isoformat()
            }
        
        print(f"Template file found: {template_path}")
        
        # Step 2: Import python-docx
        from docx import Document
        
        # Step 3: Open DOCX document
        doc = Document(template_path)
        print(f"DOCX document loaded")
        
        # Step 4: Extract structure
        sections = []
        current_section = None
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            if not text:
                continue
            
            # Identify headings
            style = para.style.name
            
            if "Heading" in style:
                # This is a section heading
                level = int(style.split()[-1]) if style[-1].isdigit() else 1
                
                section = {
                    "name": text,
                    "placeholder": text.upper().replace(" ", "_"),
                    "level": level,
                    "type": "string",
                    "required": True,
                    "character_limit": None
                }
                sections.append(section)
                current_section = text
                print(f"Found section: {text}")
            
            elif para.text.startswith("•") or para.text.startswith("-"):
                # Bullet points indicate list type
                if current_section:
                    sections[-1]["type"] = "list"
                    print(f"  → {text[:60]}...")
        
        # Step 5: Extract tables
        for table in doc.tables:
            table_rows = len(table.rows)
            table_cols = len(table.columns)
            
            table_info = {
                "name": f"Table ({table_rows}x{table_cols})",
                "placeholder": f"TABLE_{len(sections)+1}",
                "type": "table",
                "rows": table_rows,
                "columns": table_cols,
                "required": True
            }
            sections.append(table_info)
            print(f"Found table: {table_rows} rows x {table_cols} columns")
        
        # Step 6: Validate required sections
        required_keywords = ["test", "case", "acceptance", "criteria", "risk", "result"]
        found_keywords = []
        
        for section in sections:
            section_name_lower = section["name"].lower()
            for keyword in required_keywords:
                if keyword in section_name_lower and keyword not in found_keywords:
                    found_keywords.append(keyword)
        
        print(f"\n✓ Identified {len(found_keywords)} key sections")
        
        # Step 7: Save schema to .tmp/
        schema_file = ".tmp/template_schema.json"
        os.makedirs(".tmp", exist_ok=True)
        
        with open(schema_file, "w") as f:
            json.dump(sections, f, indent=2)
        
        print(f"Schema saved to: {schema_file}")
        
        return {
            "status": "valid",
            "template_name": "Test Plan - Template.docx",
            "sections": sections,
            "total_sections": len(sections),
            "message": f"Template successfully parsed. Found {len(sections)} sections.",
            "timestamp": datetime.now().isoformat()
        }
    
    except ImportError:
        return {
            "status": "error",
            "template_name": "Test Plan - Template.docx",
            "message": "python-docx library not installed. Run: pip install -r requirements.txt",
            "timestamp": datetime.now().isoformat()
        }
    
    except Exception as e:
        return {
            "status": "error",
            "template_name": "Test Plan - Template.docx",
            "message": f"Error parsing template: {str(e)}",
            "timestamp": datetime.now().isoformat()
        }

if __name__ == "__main__":
    result = parse_template()
    print("\n" + "=" * 50)
    print("📊 Parse Result:")
    print("=" * 50)
    print(json.dumps(result, indent=2))
    print("=" * 50)
