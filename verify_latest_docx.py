#!/usr/bin/env python3
from docx import Document
import os

# Find latest docx file
latest_file = None
latest_time = 0
for fname in os.listdir('.tmp'):
    if fname.endswith('.docx'):
        fpath = os.path.join('.tmp', fname)
        mtime = os.path.getmtime(fpath)
        if mtime > latest_time:
            latest_time = mtime
            latest_file = fpath

if not latest_file:
    print("No DOCX files found")
    exit(1)

print(f"Verifying: {latest_file}\n")

doc = Document(latest_file)

print("=== DOCUMENT CONTENT ===\n")
for para in doc.paragraphs:
    if para.text.strip():
        if para.style.name.startswith('Heading'):
            print(f"\n{para.style.name}: {para.text}")
        else:
            text = para.text[:80]
            print(f"  {text}")

print("\n\n=== VERIFICATION RESULTS ===")
full_text = " ".join([p.text for p in doc.paragraphs]).lower()

# Check for required sections
required_sections = [
    "test plan",
    "objective",
    "scope",
    "inclusions",
    "test environments",
    "defect reporting",
    "test strategy",
    "test schedule",
    "test deliverables"
]

print("\nRequired Sections:")
for section in required_sections:
    found = any(section in p.text.lower() for p in doc.paragraphs)
    status = "✓" if found else "✗"
    print(f"  {status} {section}")

# Check for Jira-specific content
print("\nJira-Specific Content Found:")
jira_keywords = [
    "user login",
    "credentials",
    "authentication",
    "chrome",
    "windows",
    "safari",
    "macos"
]

for keyword in jira_keywords:
    found = keyword in full_text
    status = "✓" if found else "✗"
    print(f"  {status} {keyword}")

# Check for template pollution
print("\nTemplate Content Check (should be empty):")
template_keywords = ["vwo", "react", "jquery", "postgres", "nginx", "apache", "sauce"]
template_found = [kw for kw in template_keywords if kw in full_text]

if template_found:
    print(f"  ✗ Found template keywords: {template_found}")
else:
    print(f"  ✓ No template content found")

print(f"\nTest Cases Table: ", end="")
print(f"{len(doc.tables)} table(s) found")
for t_idx, table in enumerate(doc.tables):
    print(f"\n  Table {t_idx + 1}: {len(table.rows)} rows")
    if len(table.rows) > 0:
        headers = [cell.text for cell in table.rows[0].cells]
        print(f"    Headers: {headers}")
