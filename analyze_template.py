#!/usr/bin/env python3
from docx import Document
from docx.oxml.ns import qn

doc = Document('./test_plan_templat/Test Plan - Template.docx')

print('=== TEMPLATE ANALYSIS ===\n')

# Check headers and footers
for sec_idx, section in enumerate(doc.sections):
    print(f'Section {sec_idx}:')
    
    if section.header.paragraphs:
        for para in section.header.paragraphs:
            if para.text.strip():
                print(f'  HEADER: {para.text[:80]}')
    
    if section.footer.paragraphs:
        for para in section.footer.paragraphs:
            if para.text.strip():
                print(f'  FOOTER: {para.text[:80]}')

# Search all paragraphs for the problematic text
print('\n=== SEARCHING FOR PROBLEM TEXT ===')
for idx, para in enumerate(doc.paragraphs):
    text = para.text.lower()
    if any(x in text for x in ['pramod', 'academy', 'thetest']):
        print(f'Paragraph {idx}: {para.text[:80]}')
        print(f'  Style: {para.style.name}')

# Check background/watermark in document settings
print('\n=== CHECKING DOCUMENT BACKGROUND ===')
try:
    core_props = doc.core_properties
    print(f'Title: {core_props.title}')
    print(f'Subject: {core_props.subject}')
    print(f'Author: {core_props.author}')
except:
    pass
